from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


COLOR_PRIMARY = RGBColor(0x25, 0x63, 0xEB)
COLOR_SUCCESS = RGBColor(0x10, 0xB9, 0x81)
COLOR_WARNING = RGBColor(0xF5, 0x9E, 0x0B)
COLOR_DANGER = RGBColor(0xEF, 0x44, 0x44)
COLOR_TEXT = RGBColor(0x1F, 0x29, 0x37)
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)


def _set_cell_bg(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _score_color(score: float) -> RGBColor:
    if score >= 70:
        return COLOR_SUCCESS
    if score >= 55:
        return COLOR_WARNING
    return COLOR_DANGER


def _score_label(score: float) -> str:
    if score >= 85:
        return "Excellent"
    if score >= 70:
        return "Good"
    if score >= 55:
        return "Fair"
    return "Weak"


def _format_evidence(item) -> str:
    if isinstance(item, dict):
        section = item.get("section")
        bullet_index = item.get("bullet_index")
        excerpt = item.get("excerpt") or item.get("jd_line") or str(item)
        prefix = []
        if section:
            prefix.append(str(section))
        if bullet_index:
            prefix.append(f"bullet #{bullet_index}")
        joined = " - ".join(prefix)
        return f"{joined}: {excerpt}" if joined else str(excerpt)
    return str(item)


def _add_heading(doc: Document, text: str, level: int = 2) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = COLOR_TEXT


def _add_header(doc: Document, cv_title: str, jd_title: str, generated_at: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CV Reviewer Report")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_PRIMARY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("CV to JD evaluation summary")
    sr.font.size = Pt(11)
    sr.font.color.rgb = COLOR_MUTED

    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("CV", cv_title or "N/A"),
        ("JD", jd_title or "N/A"),
        ("Generated", generated_at),
    ]
    for idx, (label, value) in enumerate(rows):
        left = table.rows[idx].cells[0]
        right = table.rows[idx].cells[1]
        _set_cell_bg(left, "EFF6FF")
        lp = left.paragraphs[0].add_run(label)
        lp.bold = True
        lp.font.color.rgb = COLOR_PRIMARY
        rp = right.paragraphs[0].add_run(value)
        rp.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()


def _add_score_summary(doc: Document, final_score: float, summary: dict) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    color = _score_color(final_score)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    score_run = p.add_run(f"{final_score:.0f}/100")
    score_run.bold = True
    score_run.font.size = Pt(28)
    score_run.font.color.rgb = color

    label_p = cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_p.add_run(summary.get("label") or _score_label(final_score))
    label_run.bold = True
    label_run.font.color.rgb = color

    if summary.get("top_priorities"):
        priority_p = cell.add_paragraph()
        priority_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = priority_p.add_run("Top priorities: " + ", ".join(summary.get("top_priorities", [])[:3]))
        pr.font.size = Pt(9)
        pr.font.color.rgb = COLOR_MUTED

    doc.add_paragraph()


def _add_axes(doc: Document, axes: dict) -> None:
    if not axes:
        return

    _add_heading(doc, "Score Axes")
    table = doc.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Overall fit", axes.get("overall_fit")),
        ("Skill coverage", axes.get("skill_coverage")),
        ("Evidence strength", axes.get("evidence_strength")),
        ("CV quality", axes.get("cv_quality")),
    ]

    for label, score in rows:
        if score is None:
            continue
        row = table.add_row().cells
        row[0].text = label
        row[1].text = "█" * int(float(score) / 5) + "░" * (20 - int(float(score) / 5))
        row[2].text = f"{float(score):.0f}/100"
        row[2].paragraphs[0].runs[0].font.color.rgb = _score_color(float(score))

    doc.add_paragraph()


def _add_breakdown(doc: Document, breakdown: dict, weights: dict) -> None:
    _add_heading(doc, "Weighted Breakdown")
    rows = [
        ("CV structure", breakdown.get("section_score", 0), weights.get("section_score")),
        ("Skill coverage", breakdown.get("skill_score", 0), weights.get("skill_score")),
        ("Semantic match", breakdown.get("semantic_score", 0), weights.get("semantic_score")),
        ("Technical keywords", breakdown.get("keyword_score", 0), weights.get("keyword_score")),
        ("Experience alignment", breakdown.get("experience_score", 0), weights.get("experience_score")),
        ("Evidence quality", breakdown.get("jd_structure_score", breakdown.get("structure_score", 0)), weights.get("jd_structure_score")),
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0].cells
    header[0].text = "Dimension"
    header[1].text = "Score"
    header[2].text = "Weight"

    for label, score, weight in rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = f"{float(score):.0f}/100"
        row[1].paragraphs[0].runs[0].font.color.rgb = _score_color(float(score))
        row[2].text = f"{float(weight):.0f}%" if weight is not None else "-"

    doc.add_paragraph()


def _add_skills_summary(doc: Document, skills_summary: dict) -> None:
    _add_heading(doc, "Skills Summary")
    coverage = float(skills_summary.get("required_coverage_pct", 0) or 0)
    cp = doc.add_paragraph()
    cr = cp.add_run(f"Required skill coverage: {coverage:.0f}%")
    cr.bold = True
    cr.font.color.rgb = _score_color(coverage)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    left, right = table.rows[0].cells
    left.paragraphs[0].add_run("Matched required").bold = True
    right.paragraphs[0].add_run("Missing required").bold = True

    matched = skills_summary.get("matched_required", []) or []
    missing = skills_summary.get("missing_required", []) or []
    max_rows = max(len(matched), len(missing), 1)
    for idx in range(max_rows):
        if idx > 0:
            row = table.add_row().cells
            left = row[0]
            right = row[1]
        if idx < len(matched):
            left.add_paragraph(str(matched[idx]))
        if idx < len(missing):
            run = right.add_paragraph(str(missing[idx])).runs[0]
            run.font.color.rgb = COLOR_DANGER

    doc.add_paragraph()


def _add_personal_info_template(cell) -> None:
    title = cell.add_paragraph()
    run = title.add_run("Personal info template")
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    for label in ["Full name", "Email", "Job", "LinkedIn", "GitHub"]:
        p = cell.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        blank = p.add_run("______________________________")
        blank.font.color.rgb = COLOR_MUTED


def _add_issues(doc: Document, issues: list[dict]) -> None:
    if not issues:
        return

    _add_heading(doc, "Issues and Recommendations")
    for issue in issues:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.rows[0].cells[0]

        severity = issue.get("severity", "low")
        fill = {"high": "FEE2E2", "medium": "FEF3C7", "low": "F3F4F6"}.get(severity, "F3F4F6")
        _set_cell_bg(cell, fill)

        title = issue.get("title") or issue.get("code", "Issue")
        p = cell.paragraphs[0]
        title_run = p.add_run(title)
        title_run.bold = True
        title_run.font.color.rgb = COLOR_TEXT

        explanation = issue.get("explanation_en") or issue.get("explanation") or ""
        if explanation:
            ep = cell.add_paragraph(explanation)
            ep.runs[0].font.color.rgb = COLOR_TEXT

        evidence = issue.get("evidence", []) or issue.get("details", []) or []
        if evidence:
            cell.add_paragraph("Evidence").runs[0].bold = True
            for item in evidence[:4]:
                cell.add_paragraph(_format_evidence(item), style="List Bullet")

        if issue.get("code") == "contact_info":
            _add_personal_info_template(cell)
        else:
            suggestion = issue.get("suggested_fix_en") or issue.get("suggested_fix") or ""
            if suggestion:
                sp = cell.add_paragraph()
                sp.add_run("Suggested improvement: ").bold = True
                sp.add_run(suggestion)

        doc.add_paragraph()


def _add_section_analysis(doc: Document, section_analysis: dict) -> None:
    if not section_analysis:
        return

    _add_heading(doc, "CV Structure Analysis")
    found = section_analysis.get("sections_found", []) or []
    missing = section_analysis.get("missing_required_sections", []) or []

    fp = doc.add_paragraph()
    fp.add_run("Sections found: ").bold = True
    fp.add_run(", ".join(found) if found else "None")

    mp = doc.add_paragraph()
    mp.add_run("Missing required sections: ").bold = True
    miss_run = mp.add_run(", ".join(missing) if missing else "None")
    miss_run.font.color.rgb = COLOR_DANGER if missing else COLOR_SUCCESS

    doc.add_paragraph()


def _add_semantic_analysis(doc: Document, semantic_analysis: dict) -> None:
    if not semantic_analysis:
        return

    _add_heading(doc, "Evidence Coverage")
    status = semantic_analysis.get("status", "unavailable")
    score = float(semantic_analysis.get("score", 0) or 0)

    sp = doc.add_paragraph()
    sp.add_run("Semantic status: ").bold = True
    sp.add_run(status)

    score_p = doc.add_paragraph()
    score_p.add_run("Semantic score: ").bold = True
    score_run = score_p.add_run(f"{score:.0f}/100")
    score_run.font.color.rgb = _score_color(score)

    unmatched = semantic_analysis.get("unmatched_jd_lines", []) or []
    if unmatched:
        doc.add_paragraph("JD lines that still need evidence").runs[0].bold = True
        for item in unmatched[:5]:
            doc.add_paragraph(_format_evidence(item), style="List Bullet")

    doc.add_paragraph()


def _add_footer(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Generated by CV Reviewer. This report is advisory and should be reviewed before final submission."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED


def generate_match_report_docx(match_record, report_json: dict) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    summary = report_json.get("summary", {})
    breakdown = report_json.get("score_breakdown", {})
    axes = report_json.get("score_axes", {})
    weights = report_json.get("score_weights", {})
    skills_summary = report_json.get("skills_summary", {})
    issues = report_json.get("issues", [])
    section_analysis = report_json.get("section_analysis", {})
    semantic_analysis = report_json.get("semantic_analysis", {})

    final_score = float(summary.get("final_score", 0) or 0)
    generated_at = (
        match_record.created_at.strftime("%B %d, %Y %H:%M")
        if getattr(match_record, "created_at", None)
        else datetime.now().strftime("%B %d, %Y %H:%M")
    )

    _add_header(
        doc,
        summary.get("cv_title", ""),
        summary.get("jd_title", ""),
        generated_at,
    )
    _add_score_summary(doc, final_score, summary)
    _add_axes(doc, axes)
    _add_breakdown(doc, breakdown, weights)
    _add_skills_summary(doc, skills_summary)
    _add_section_analysis(doc, section_analysis)
    _add_semantic_analysis(doc, semantic_analysis)
    _add_issues(doc, issues)
    _add_footer(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
